import requests, bs4
import openpyxl
import docx, os
import sqlite3


def DoIt(actor):

    try:
        # Establish a connection with the sqlite database file
        conn = sqlite3.connect("database.sqlite")
        cur = conn.cursor()

        # Creating the table if it dosen't exist in the database 
        cur.execute("CREATE TABLE IF NOT EXISTS ACTORHUNT (NAME TEXT PRIMARY KEY, DOB TEXT, JOB TEXT, PICTURE TEXT, INFO TEXT)")
        cur.execute("CREATE TABLE IF NOT EXISTS MOVIES (ACTORNAME TEXT, TITLE TEXT, YEAR TEXT, FOREIGN KEY (ACTORNAME) REFERENCES ACTORHUNT (NAME) )")

        # SQLite query to search for an actor in the table ACTORHUNT of database
        cur.execute("SELECT * FROM ACTORHUNT WHERE NAME LIKE ?",('%'+actor.lower()+'%',))
        data0 = cur.fetchone() # retrieving all details from the returned search cursor

        # SQLite query to search for all the movies of the actor in the table MOVIES of database
        cur.execute("SELECT * FROM MOVIES WHERE ACTORNAME LIKE ?",('%'+actor.lower()+'%',))  
        data1 = cur.fetchall()
    except sqlite3.Error as e:
        print("database error:\n" ,e.args)

    try:
        # If the data is found in the database then close connection with database
        # write the data into respective Doc and Excel file
        # return the required data to the client 
        if data0 is not None:
            cur.close()
            conn.close()
            writeDoc(data0)
            writeXL(data1)
            return data0
        else:
            # If the data not found in the database then close connection with database
            # call `fetchAndSave(actor)` function to scrape web for the details of the actor
            cur.close()
            conn.close()
            return fetchAndSave(actor)
    except UnboundLocalError as e:
        print("Error in database so \n" ,e.args)


def fetchAndSave(actor):
    try:
        # making request to the imdb website with the search query of the actor
        res = requests.get("https://www.imdb.com/find?q=" + actor.replace("", "+"))
        res.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(e)

    # parsing html with BeautifulSoup
    soup = bs4.BeautifulSoup(res.text, "html.parser")
    soupelem = soup.select(".result_text a")

    try:
        actorpage = requests.get("https://www.imdb.com" + soupelem[0].get("href"))
        actorpage.raise_for_status()
    except IndexError as e:
        print(e)
        return ['NOT FOUND', 'nil', 'nil', 'nil', 'nil']

    # actor page 
    try:
        soup = bs4.BeautifulSoup(actorpage.text, "html.parser")

        # actor info
        actorName = soup.select(".header .itemprop")[0].text
        actorJobs = list(
            map(lambda name: name.text[1:], soup.select("#name-job-categories a span"))
        )
        actorPhoto = soup.select("#name-poster")[0].get("src")

        actorBDsoup = soup.select("#name-born-info")[0]
        actorBD = " ".join(
            list(map(lambda text: text.strip(), actorBDsoup.text.split("\n")))
        )

        # all the movies of the actor
        soupelem = soup.select("#filmo-head-actor + .filmo-category-section .filmo-row")
        if len(soupelem) == 0:
            soupelem = soup.select("#filmo-head-actress + .filmo-category-section .filmo-row")
        fullBioLink = soup.find("span", {"class":"see-more inline nobr-only"}).a.get("href")
    except IndexError as e:
        print("actor not found, try again")
        return ['NOT FOUND', 'nil', 'nil', 'nil', 'nil']

    try:
        bio = requests.get("https://www.imdb.com"+fullBioLink)  
        bio.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(e)

    bioSoup = bs4.BeautifulSoup(bio.text, "html.parser")
    actorData = bioSoup.find("div", {"class":"soda odd"}).p

# -------------------------------------------SAVING TO THE DATABASE---------------------------------------
    try:
        conn = sqlite3.connect("database.sqlite")
        cur = conn.cursor()

        print("writing database")

        cur.execute("INSERT INTO ACTORHUNT (NAME, DOB, JOB, PICTURE, INFO) VALUES (?, ?, ?, ?, ?)",(str(actorName).lower(), str(actorBD), ', '.join(actorJobs), str(actorPhoto), str(actorData.text),))
        conn.commit()

        for row in range(2, len(soupelem)+2):
            movie = soupelem[row-2]
            title = movie.select("b a")[0].text
            year = movie.select(".year_column")[0].text[:6]
            cur.execute("INSERT INTO MOVIES (ACTORNAME, TITLE, YEAR) VALUES (?, ?, ?)",(str(actorName).lower(), str(title), str(year)))
            conn.commit()

        cur.execute("SELECT * FROM ACTORHUNT WHERE NAME = ?",(actorName.lower(),))
        data0 = cur.fetchone()

        cur.execute("SELECT * FROM MOVIES WHERE ACTORNAME = ?",(actorName.lower(),))  
        data1 = cur.fetchall()

        cur.close()
        conn.close()
    except sqlite3.Error as e:
        print("database key constraint error:\n" , e.args)
        # if key constrint occurs during insertion then the client is prompted to check the spelling of the search query
        return ['check spelling', 'nil', 'nil', 'nil', 'nil'] 
# -------------------------------------------------------------------------------------------------------------

    try:
        writeDoc(data0)
    except Exception as e:
        print(e)

    try:
        writeXL(data1)
    except Exception as e:
        print(e)

    try:
        return data0
    except UnboundLocalError as e:
        print("database error so \n" ,e.args)
        return ["NOT FOUND", 'NIL', 'NIL', 'NIL', 'NIL']


# function to write the actor movies into the excel sheet with extension .xlsx
def writeXL(data1):
    wb = openpyxl.Workbook()
    sheet = wb.get_sheet_by_name("Sheet")
    sheet["A1"].value = "Title"
    sheet["B1"].value = "Year"

    for row in range(len(data1)):
        sheet["A" + str(row+2)] = data1[row][1]
        sheet["B" + str(row+2)] = data1[row][2]

    wb.save("static/actorMovies.xlsx")
    print("xlsx success")

# function to write the actor data into the Word document with extension .docx
def writeDoc(data0):
    os.makedirs("photos", exist_ok=True)
    try:
        res = requests.get(data0[3])
    except requests.exceptions.ConnectionError as e:
        print(e)

    imageFile = open(os.path.join("photos", os.path.basename(data0[3])), "wb")
    actorPhotoPath = data0[3].split("/")[-1]


    for chunk in res.iter_content(100000):
        imageFile.write(chunk)
    imageFile.close()

    doc = docx.Document()
    doc.add_paragraph(data0[0].capitalize(), "Title")
    doc.add_picture("./photos/" + actorPhotoPath, width=docx.shared.Cm(5))
    doc.add_paragraph("Actor Jobs: " + data0[2])
    doc.add_paragraph(data0[1])
    doc.add_paragraph(data0[4])

    doc.save("static/actordesc.docx")
    print("docx success!")
